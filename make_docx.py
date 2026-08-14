"""Render the Methods/Results/Discussion markdown into a .docx manuscript draft.

Neither pandoc nor docx-js is available in this environment, so the conversion is
done directly with python-docx.  Markdown tables become real Word tables and
display equations become centred, monospaced LaTeX in a shaded block — Word has no
way to typeset LaTeX natively, and silently mangling the maths would be worse than
presenting it as source that can be pasted into Word's equation editor (Insert >
Equation > paste LaTeX) or straight into a LaTeX manuscript.
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

SRC = "outputs/PAPER_Methods_Results_Discussion.md"
DST = "outputs/PAPER_Methods_Results_Discussion.docx"

ACCENT = RGBColor(0xD5, 0x51, 0x81)
INK = RGBColor(0x0B, 0x0B, 0x0B)
MUTED = RGBColor(0x52, 0x51, 0x4E)
MATH_INK = RGBColor(0x18, 0x4F, 0x95)
MATH_FONT = "DejaVu Sans Mono"
BODY_FONT = "Arial"


def shade(cell_or_par, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    target = cell_or_par._tc.get_or_add_tcPr() if hasattr(cell_or_par, "_tc") \
        else cell_or_par._p.get_or_add_pPr()
    target.append(el)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\$[^$]+?\$)")


UNESCAPE = re.compile(r"\\([*_~`\[\]()#+\-.!])")


def add_runs(par, text: str) -> None:
    """Render bold / italic / code / inline-maths spans."""
    for part in INLINE.split(text):
        if not part:
            continue
        if not part.startswith("$"):
            part = UNESCAPE.sub(r"\1", part)
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2])
            r.bold = True
            r.font.name = BODY_FONT
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = MUTED
        elif part.startswith("$") and part.endswith("$"):
            r = par.add_run(part[1:-1])
            r.font.name = MATH_FONT
            r.font.size = Pt(9.5)
            r.font.color.rgb = MATH_INK
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = par.add_run(part[1:-1])
            r.italic = True
            r.font.name = BODY_FONT
        else:
            par.add_run(part).font.name = BODY_FONT


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def split_row(line: str):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main() -> None:
    lines = open(SRC, encoding="utf-8").read().split("\n")
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(0.9)

    i, n_tab, n_eq = 0, 0, 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "$$":
            body, i = [], i + 1
            while i < len(lines) and lines[i].strip() != "$$":
                body.append(lines[i].rstrip())
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            txt = " ".join(x.strip() for x in body if x.strip())
            txt = txt.replace("\\boxed{", "").rstrip("}") if "\\boxed" in txt else txt
            r = p.add_run(txt)
            r.font.name = MATH_FONT
            r.font.size = Pt(9.5)
            r.font.color.rgb = MATH_INK
            shade(p, "F4F6FA")
            n_eq += 1
            continue

        if is_table_row(line) and i + 1 < len(lines) and set(
                lines[i + 1].strip().strip("|").replace(":", "").replace("-", "")
                .replace("|", "").strip()) == set():
            header = split_row(line)
            i += 2
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(split_row(lines[i]))
                i += 1
            ncol = len(header)
            t = doc.add_table(rows=1, cols=ncol)
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            avail = 6.7
            widths = [Inches(avail / ncol)] * ncol
            for c, h in enumerate(header):
                cell = t.rows[0].cells[c]
                cell.width = widths[c]
                cell.text = ""
                add_runs(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
                    r.font.size = Pt(9)
                shade(cell, "EAEFF7")
            for row in rows:
                cells = t.add_row().cells
                for c in range(ncol):
                    cells[c].width = widths[c]
                    cells[c].text = ""
                    add_runs(cells[c].paragraphs[0], row[c] if c < len(row) else "")
                    for r in cells[c].paragraphs[0].runs:
                        r.font.size = Pt(9)
            doc.add_paragraph()
            n_tab += 1
            continue

        if stripped.startswith("#"):
            lvl = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[lvl:].strip()
            h = doc.add_heading(level=min(lvl, 4))
            add_runs(h, text)
            for r in h.runs:
                r.font.color.rgb = ACCENT if lvl <= 2 else INK
            i += 1
            continue

        if stripped in {"---", "***"}:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            b = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "C9C9C4")
            b.append(bot)
            pPr.append(b)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            body = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            i += 1
            for b in body:
                p = doc.add_paragraph()
                r = p.add_run(b)
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                shade(p, "F2F2EF")
            continue

        m = re.match(r"^(\s*)([*\-])\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3 + 0.25 * (len(m.group(1)) // 2))
            add_runs(p, m.group(3))
            i += 1
            continue
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(3))
            i += 1
            continue

        if not stripped:
            i += 1
            continue
        buf = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(
                ("#", "|", "$$", "```", "---", "* ", "- ")) and not re.match(
                r"^\s*\d+\.\s", lines[i]):
            buf.append(lines[i].strip())
            i += 1
        add_runs(doc.add_paragraph(), " ".join(buf))

    doc.save(DST)
    print(f"wrote {DST}: {n_tab} tables, {n_eq} equations, "
          f"{len(doc.paragraphs)} paragraphs")


if __name__ == "__main__":
    sys.exit(main())
